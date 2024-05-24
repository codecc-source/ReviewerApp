
import random
import tkinter as tk
import json
import docx
from docx.shared import Inches
from tkinter import simpledialog, messagebox, filedialog
from PIL import Image, ImageTk
from tkinter import ttk
from gtts import gTTS


class NotesReviewer:
    def __init__(self, root):
        self.image_preview_window = None
        self.image_label = None
        self.root = root
        self.root.title("Notes Reviewer")
        self.notes = {}
        self.load_notes()

        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        root.geometry(f"{screen_width}x{screen_height}")

        root.attributes('-fullscreen', False)

        self.menu_frame = tk.Frame(root, bg="#36454f")
        self.menu_frame.pack_propagate(False)
        self.menu_frame.pack(fill=tk.BOTH, expand=True)

        label_font = ("Arial", 10)
        button_font = ("Arial", 10, "bold")

        self.question_label = tk.Label(self.menu_frame, text="Enter Question(text or image):", bg='#555D50', fg='white', font=20, width=1080)
        self.question_label.pack()

        self.question_text_entry = tk.Text(self.menu_frame, width=80, height=3, wrap=tk.WORD, bg='lightblue')
        self.question_text_entry.pack(pady=5)

        self.browse_image_button = tk.Button(self.menu_frame, text="Browse Image(for visual questions)", command=self.browse_image, font=button_font, bg='#EA3B52', fg='white', width=50)
        self.browse_image_button.pack(pady=5)

        self.answer_label = tk.Label(self.menu_frame, text="Enter Answer:", font=20, bg='#555D50', fg='white', width=1080)
        self.answer_label.pack()

        self.answer_entry = tk.Text(self.menu_frame, width=80, height=2, wrap=tk.WORD, bg='lightblue')
        self.answer_entry.pack(pady=5)

        self.add_note_button = tk.Button(self.menu_frame, text="Add Note", command=self.add_note, font=button_font, width=50, bg='#666362', fg='white')
        self.add_note_button.pack()

        self.review_button = tk.Button(self.menu_frame, text="Review Notes(Quiz)", command=self.review_notes, font=button_font, width=50, bg='#666362', fg='white')
        self.review_button.pack()

        self.clear_button = tk.Button(self.menu_frame, text="Clear Notes", command=self.clear_notes, font=button_font, width=50, bg='#666362', fg='white')
        self.clear_button.pack()

        self.undo_button = tk.Button(self.menu_frame, text="Undo", command=self.undo_last_note, font=button_font, width=50, bg='#666362', fg='white')
        self.undo_button.pack()

        self.edit_note_button = tk.Button(self.menu_frame, text="Edit Note", command=self.edit_note, font=button_font, width=50, bg='#666362', fg='white')
        self.edit_note_button.pack()

        self.export_button = tk.Button(self.menu_frame, text="Export to Word", command=self.export_to_word, font=button_font, width=50, bg='#666362', fg='white')
        self.export_button.pack()

        self.show_notes_button = tk.Button(self.menu_frame, text="Show Notes List", command=self.show_notes_list, font=button_font, width=50, bg='#666362', fg='white')
        self.show_notes_button.pack()

        self.json_filename_label = tk.Label(self.menu_frame, text="Note File Name \nEx: Subject1.json):", font=label_font, width=50, bg='#3B3131', fg='white')
        self.json_filename_label.pack(pady=5)
        self.json_filename_entry = tk.Entry(self.menu_frame, justify='center', width=105, bg='lightblue')
        self.json_filename_entry.pack()
        self.json_filename_entry.insert(0, "SubjectName.json")

        self.save_notes_button = tk.Button(self.menu_frame, text="Save Notes to JSON", command=self.save_notes_to_json, font=button_font, width=50, bg='#666362', fg='white')
        self.save_notes_button.pack()

        self.load_json_button = tk.Button(self.menu_frame, text="Load Notes from JSON", command=self.load_notes_from_json, font=button_font, width=50, bg='#666362', fg='white')
        self.load_json_button.pack()

        self.highest_score_label = tk.Label(self.menu_frame, text="Highest Score(Per Session): 0", font=label_font, width=50, bg='#3B3131', fg='white')
        self.highest_score_label.pack()

        self.save_mp3_button = tk.Button(self.menu_frame, text="Save Notes as MP3", command=self.save_notes_as_mp3, font=button_font, width=50, bg='#666362', fg='white')
        self.save_mp3_button.pack()

        self.instructions_button = tk.Button(self.menu_frame, text="Instructions", command=self.show_instructions, font=button_font, width=50, bg='#666362', fg='white')
        self.instructions_button.pack()

        self.quit_button = tk.Button(self.menu_frame, text="Quit", command=root.quit, font=button_font, width=1080, bg='#DA2C43', fg='white')
        self.quit_button.pack(pady=5)

        self.image_path = None
        self.highest_score = 0

    def browse_image(self):
        file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg *.gif *.bmp")])
        if file_path:
            self.image_path = file_path
            self.load_image(file_path)

    def load_image(self, file_path):
        try:
            image = Image.open(file_path)
            image.thumbnail((700, 700))
            photo = ImageTk.PhotoImage(image)
            self.image_label = tk.Label(self.menu_frame, image=photo)
            self.image_label.image = photo

            self.image_preview_window = tk.Toplevel(self.root)
            self.image_preview_window.title("Notes Image Review")

            screen_width = self.image_preview_window.winfo_screenwidth()
            screen_height = self.image_preview_window.winfo_screenheight()

            window_width = image.width
            window_height = image.height

            x = (screen_width - window_width) // 2
            y = int(0.7 * (screen_height - window_height))

            self.image_preview_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

            image_preview_label = tk.Label(self.image_preview_window, image=photo)
            image_preview_label.image = photo
            image_preview_label.pack()
            close_button = tk.Button(self.image_preview_window, text="Close", command=self.close_image_preview)
            close_button.pack()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load image: {str(e)}")
            self.image_path = None

    def close_image_preview(self):
        self.image_preview_window.destroy()

    def add_note(self):
        answer = self.answer_entry.get("1.0", tk.END).strip()

        if not answer:
            messagebox.showerror("Error", "Please enter an answer before adding a note.")
            return

        if self.image_path:
            question = f"[Image] {self.image_path}"
            self.image_path = None
            self.image_label.destroy()
            self.close_image_preview()
        else:
            question = self.question_text_entry.get("1.0", tk.END).strip()

        if question or self.image_path:
            if question not in self.notes:
                self.notes[question] = answer
                messagebox.showinfo("Note Added", "Note added successfully!")
                self.question_text_entry.delete("1.0", tk.END)
                self.answer_entry.delete("1.0", tk.END)
            else:
                messagebox.showerror("Duplicate Note", "This note already exists.")
        else:
            messagebox.showerror("Error", "Please enter either a question or an image as a question.")

    def review_notes(self):
        if not self.notes:
            messagebox.showinfo("No Notes", "No notes available.")
            return

        correct_count = 0
        total_questions = len(self.notes)

        questions = list(self.notes.keys())
        random.shuffle(questions)

        correct_answers_message = "Correct Answers:\n\n"
        incorrect_answers_message = "Incorrect Answers:\n\n"

        for question in questions:
            review_window = tk.Toplevel(self.root)
            review_window.grab_set()
            review_window.title("Review Note")

            if question.startswith("[Image]"):
                image_path = question.replace("[Image] ", "")
                image = Image.open(image_path)
                image.thumbnail((380, 380))
                photo = ImageTk.PhotoImage(image)
                image_label = tk.Label(review_window, image=photo)
                image_label.image = photo
                image_label.pack()

                user_answer = simpledialog.askstring("Review Note", "See the image below and answer.")
                image_label.destroy()
            else:
                user_answer = simpledialog.askstring("Review Note", f"Q: {question}\nYour Answer:")

            if user_answer is None:
                review_window.destroy()
                return
            elif user_answer.strip() == "":
                messagebox.showinfo("Incomplete Answer", "Please provide an answer for the question.")
                review_window.destroy()
                return

            user_words = set(user_answer.lower().split())
            stored_words = set(self.notes[question].lower().split())

            if user_words.issubset(stored_words) or stored_words.issubset(user_words):
                correct_count += 1
                correct_answers_message += f"Q: {question}\nA: {self.notes[question]}\n\n"
            else:
                incorrect_answers_message += f"Q: {question}\nYour Answer: {user_answer}\nCorrect Answer: {self.notes[question]}\n\n"

            review_window.destroy()

        review_message = f"You got {correct_count}/{total_questions} correct."

        if correct_count > self.highest_score:
            self.highest_score = correct_count
            self.highest_score_label.config(text=f"Highest Score(Per Session): {self.highest_score}")

        if correct_count > 0:
            messagebox.showinfo("Score", review_message + "\n\n" + correct_answers_message)
        else:
            messagebox.showinfo("Retry", review_message + "\n\n" + correct_answers_message)

        if incorrect_answers_message != "Incorrect Answers:\n\n":
            messagebox.showinfo("Incorrect Answers", incorrect_answers_message)

    def clear_notes(self):
        confirmation = messagebox.askyesno("Clear Notes", "Are you sure you want to clear all notes?")
        if confirmation:
            self.notes = {}
            self.highest_score_label.config(text="Highest Score(Per Session): 0")

    def undo_last_note(self):
        if not self.notes:
            messagebox.showinfo("No Notes", "No notes available.")
            return

        last_question = list(self.notes.keys())[-1]
        del self.notes[last_question]
        self.save_notes()
        messagebox.showinfo("Undo", "Last note removed successfully.")

    def edit_note(self):
        if not self.notes:
            messagebox.showinfo("No Notes", "No notes available.")
            return

        available_notes = list(self.notes.keys())

        edit_window = tk.Toplevel(self.root)
        edit_window.title("Edit Note")

        label_font = ("Arial", 14)

        selected_note_label = tk.Label(edit_window, text="Select a note to edit:", font=label_font)
        selected_note_label.pack()

        selected_note_combobox = ttk.Combobox(edit_window, values=available_notes)
        selected_note_combobox.set(available_notes[0])
        selected_note_combobox.pack()

        question_label = tk.Label(edit_window, text="Question:", font=label_font)
        question_label.pack()
        question_text_entry = tk.Text(edit_window, width=80, height=3, wrap=tk.WORD)
        question_text_entry.pack()

        answer_label = tk.Label(edit_window, text="Answer:", font=label_font)
        answer_label.pack()
        answer_entry = tk.Text(edit_window, width=80, height=2, wrap=tk.WORD)
        answer_entry.pack()

        def update_note():
            selected_note = selected_note_combobox.get()
            new_question = question_text_entry.get("1.0", tk.END).strip()
            new_answer = answer_entry.get("1.0", tk.END).strip()

            if new_question and new_answer:
                if selected_note != new_question:
                    del self.notes[selected_note]
                self.notes[new_question] = new_answer
                self.save_notes()
                messagebox.showinfo("Note Edited", "Note edited successfully.")
                edit_window.destroy()
            else:
                messagebox.showerror("Error", "Both question and answer must be provided.")

        def browse_new_image():
            file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg *.gif *.bmp")])
            if file_path:
                question_text_entry.delete("1.0", tk.END)
                question_text_entry.insert(tk.END, f"[Image] {file_path}")

                try:
                    image = Image.open(file_path)
                    image.thumbnail((300, 300))
                    photo = ImageTk.PhotoImage(image)
                    image_preview.config(image=photo)
                    image_preview.image = photo
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to load image: {str(e)}")
                    image_preview.config(image=None)

        update_button = tk.Button(edit_window, text="Update Note", command=update_note, font=label_font)
        update_button.pack(pady=15)

        browse_image_button = tk.Button(edit_window, text="Browse New Image", command=browse_new_image,
                                         font=label_font)
        browse_image_button.pack(pady=5)

        def cancel_edit():
            edit_window.destroy()

        cancel_button = tk.Button(edit_window, text="Cancel", command=cancel_edit, font=label_font)
        cancel_button.pack(pady=5)

        image_preview_label = tk.Label(edit_window, text="Image Preview:", font=label_font)
        image_preview_label.pack()

        image_preview = tk.Label(edit_window)
        image_preview.pack()

        def select_note_to_edit(event):
            selected_note = selected_note_combobox.get()
            if selected_note.startswith("[Image]"):
                question_text_entry.delete("1.0", tk.END)
                question_text_entry.insert(tk.END, selected_note)
                browse_image_button.config(state=tk.NORMAL)

                image_path = selected_note.replace("[Image] ", "")
                try:
                    image = Image.open(image_path)
                    image.thumbnail((200, 200))
                    photo = ImageTk.PhotoImage(image)
                    image_preview.config(image=photo)
                    image_preview.image = photo
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to load image: {str(e)}")
                    image_preview.config(image=None)
            else:
                question_text_entry.delete("1.0", tk.END)
                question_text_entry.insert(tk.END, selected_note)
                browse_image_button.config(state=tk.DISABLED)
                image_preview.config(image=None)

        selected_note_combobox.bind("<<ComboboxSelected>>", select_note_to_edit)

    def export_to_word(self):
        if not self.notes:
            messagebox.showinfo("No Notes", "No notes available to export.")
            return

        doc = docx.Document()
        doc.add_heading("Notes", 0)

        for question, answer in self.notes.items():
            table = doc.add_table(rows=1, cols=2)
            table.columns[0].width = Inches(3)

            if question.startswith("[Image]"):
                image_path = question.replace("[Image] ", "")
                image_cell = table.cell(0, 0)
                image_cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(3))
            else:
                question_cell = table.cell(0, 0)
                question_cell.paragraphs[0].add_run("Question:").bold = True
                question_paragraph = question_cell.add_paragraph()
                question_paragraph.add_run(question)

            answer_cell = table.cell(0, 1)
            answer_cell.paragraphs[0].add_run("Answer:").bold = True
            answer_paragraph = answer_cell.add_paragraph()
            answer_paragraph.add_run(answer)
            doc.add_paragraph(
                "_________________________________________________________________________________________________________")

        file_name = simpledialog.askstring("Save as Word Document", "Enter a file name:")
        if file_name:
            doc.save(f"{file_name}.docx")
            messagebox.showinfo("Export Complete", f"Notes exported to {file_name}.docx")

    def show_notes_list(self):
        if not self.notes:
            messagebox.showinfo("Notes List", "No notes available.")
            return

        notes_window = tk.Toplevel(self.root)
        notes_window.title("Notes List")

        frame = tk.Frame(notes_window, width=700, height=700)
        frame.pack(fill=tk.BOTH, expand=True)

        canvas = tk.Canvas(frame, width=700, height=700)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = tk.Scrollbar(frame, orient=tk.VERTICAL, command=canvas.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        canvas.configure(yscrollcommand=scrollbar.set)

        second_frame = tk.Frame(canvas)
        canvas.create_window((0, 0), window=second_frame, anchor=tk.NW)

        for i, (question, answer) in enumerate(self.notes.items()):
            note_label = tk.Label(second_frame, text=f"Question: {question}\nAnswer: {answer}", wraplength=600, justify=tk.LEFT)
            note_label.grid(row=i, column=0, sticky=tk.W, padx=10, pady=10)

            if question.startswith("[Image]"):
                image_path = question.replace("[Image] ", "")
                image = Image.open(image_path)
                image.thumbnail((250, 250))
                photo = ImageTk.PhotoImage(image)

                image_label = tk.Label(second_frame, image=photo)
                image_label.image = photo
                image_label.grid(row=i, column=1, padx=10, pady=10)

        second_frame.update_idletasks()

        canvas.config(scrollregion=canvas.bbox(tk.ALL))

    def save_notes_to_json(self):
        json_filename = self.json_filename_entry.get()

        if not json_filename.endswith(".json"):
            messagebox.showerror("Error", "Please enter a valid JSON file name ending with .json.")
            return

        formatted_notes = []
        for question, answer in self.notes.items():
            formatted_notes.append({
                "question": question,
                "answer": answer
            })

        with open(json_filename, "w") as file:
            json.dump(formatted_notes, file, indent=2)

        messagebox.showinfo("Notes Saved", f"Notes saved to {json_filename}.")

    def save_notes(self):
        json_filename = self.json_filename_entry.get()
        if not json_filename:
            json_filename = "SubjectName.json"

        with open(json_filename, "w") as file:
            json.dump(self.notes, file)

    def load_notes(self):
        try:
            with open("notes.json", "r") as file:
                self.notes = json.load(file)
        except FileNotFoundError:
            self.notes = {}

    def load_notes_from_json(self):
        json_filename = filedialog.askopenfilename(filetypes=[("JSON Files", "*.json")])
        if json_filename:
            try:
                with open(json_filename, "r") as file:
                    data = json.load(file)
                    self.notes = {}
                    for entry in data:
                        question = entry.get("question")
                        answer = entry.get("answer")
                        if question and answer:
                            self.notes[question] = answer
                    messagebox.showinfo("Notes Loaded", f"Notes loaded from {json_filename}.")
            except FileNotFoundError:
                messagebox.showerror("Error", f"File not found: {json_filename}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load notes: {str(e)}")

    def save_notes_as_mp3(self):
        if not self.notes:
            messagebox.showinfo("No Notes", "No notes available to save as MP3.")
            return

        text_to_speech = ""
        for question, answer in self.notes.items():
            if not question.startswith("[Image]"):
                text_to_speech += f"Question: {question}\nAnswer: {answer}\n"

        if not text_to_speech:
            messagebox.showinfo("No Text Notes", "No text notes available to save as MP3.")
            return

        mp3_filename = filedialog.asksaveasfilename(defaultextension=".mp3", filetypes=[("MP3 Files", "*.mp3")])
        if mp3_filename:
            try:
                tts = gTTS(text=text_to_speech, lang="en")
                tts.save(mp3_filename)
                messagebox.showinfo("MP3 Saved", f"Notes saved as {mp3_filename}.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save as MP3: {str(e)}")

    def show_instructions(self):
        instructions_text = """
        Instructions: \n
            1. Enter Question then Enter Answer 
            * For image type questions click browse image then type answer(no need to type a question) *\n
            2. Click Add Note button
            * You can add more notes once note added confirmation is seen *\n
            3. Review Notes for a quizlet of the notes added\n
            4. Clear Notes to start over\n
            5. Undo last note entered\n
            6. Edit note to edit a particular note in the list\n
            7. Export to word. Enter file name and find the DOCX file in the folder of this program\n
            8. Show Notes List to check notes entered\n\n
            *****Next part is important if you want to keep reviewing the notes and be able to upload the notes on the quizlet website*****\n\n
            9. Enter note title for example: "Subject_A_Prelim_Notes.json"
            * Please note that the .json extension is needed so you need to type it after the note title *\n
            10. Click Save Notes to JSON once the above step is done(JSON notes will be stored in the same folder as this program)\n
            11. Load Note from JSON if you have made previous notes in JSON format(JSON notes will be stored in the same folder as this program)\n
            12. High Score for fun and challenge\n
            13. Save Notes as MP3 to listen to your notes on the go(jeepney rides for example)\n"""

        instructions_popup = tk.Toplevel(self.root)
        instructions_popup.title("Instructions")

        instructions_label = tk.Label(instructions_popup, text=instructions_text, justify=tk.LEFT)
        instructions_label.pack(padx=20, pady=20)


if __name__ == "__main__":
    root = tk.Tk()
    app = NotesReviewer(root)
    root.lift()
    root.mainloop()